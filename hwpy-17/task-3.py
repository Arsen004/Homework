'''Домашнее задание №11: Работа с комплексными файлами - excel, json,
word. Git и Jira'''

from docx import Document
from docx.shared import Pt

def create_word_file():

    doc = Document()
    doc.add_paragraph('Hello Python')


    doc.save('hello_python.docx')

def read_bold_text():

    doc = Document('hello_python.docx')


    bold_text = ''

    for para in doc.paragraphs:
        for run in para.runs:
            if run.bold:
                bold_text += run.text

    print("Bold Text:", bold_text)

def create_custom_paragraph():

    doc = Document()

    para = doc.add_paragraph('This is a new paragraph with custom font and size.')

    run = para.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(14)

    doc.save('custom_paragraph.docx')

if __name__ == "__main__":
    create_word_file()
    read_bold_text()
    create_custom_paragraph()
